"""Tests for the resume text extraction module."""

from io import BytesIO

import fitz
from docx import Document

from modules.resume_parser import extract_resume_text


def test_valid_pdf_extraction():
    """Text must be extracted from a valid PDF."""

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "John Doe Python Developer")
    pdf_content = document.tobytes()
    document.close()

    text, success, error = extract_resume_text(
        "resume.pdf",
        pdf_content,
    )

    assert success is True
    assert "John Doe" in text
    assert error == ""


def test_valid_docx_extraction():
    """Text must be extracted from a valid DOCX file."""

    document = Document()
    document.add_paragraph("Jane Doe Data Analyst")

    file_stream = BytesIO()
    document.save(file_stream)

    text, success, error = extract_resume_text(
        "resume.docx",
        file_stream.getvalue(),
    )

    assert success is True
    assert "Jane Doe" in text
    assert error == ""


def test_empty_pdf():
    """An empty PDF must fail safely."""

    text, success, error = extract_resume_text(
        "empty.pdf",
        b"",
    )

    assert text == ""
    assert success is False
    assert error == "The PDF file is empty."


def test_empty_docx():
    """An empty DOCX file must fail safely."""

    text, success, error = extract_resume_text(
        "empty.docx",
        b"",
    )

    assert text == ""
    assert success is False
    assert error == "The DOCX file is empty."


def test_corrupted_pdf():
    """A corrupted PDF must fail without crashing."""

    text, success, error = extract_resume_text(
        "corrupted.pdf",
        b"This is not a real PDF.",
    )

    assert text == ""
    assert success is False
    assert "The PDF could not be processed:" in error


def test_corrupted_docx():
    """A corrupted DOCX file must fail without crashing."""

    text, success, error = extract_resume_text(
        "corrupted.docx",
        b"This is not a real DOCX file.",
    )

    assert text == ""
    assert success is False
    assert "The DOCX file could not be processed:" in error


def test_textless_pdf():
    """A PDF without readable text must fail safely."""

    document = fitz.open()
    document.new_page()
    pdf_content = document.tobytes()
    document.close()

    text, success, error = extract_resume_text(
        "textless.pdf",
        pdf_content,
    )

    assert text == ""
    assert success is False
    assert error == "No readable text was found in the PDF."


def test_empty_docx_document():
    """A valid DOCX without text must fail safely."""

    document = Document()

    file_stream = BytesIO()
    document.save(file_stream)

    text, success, error = extract_resume_text(
        "empty_document.docx",
        file_stream.getvalue(),
    )

    assert text == ""
    assert success is False
    assert error == "No readable text was found in the DOCX file."


def test_unsupported_format():
    """An unsupported file type must fail safely."""

    text, success, error = extract_resume_text(
        "resume.txt",
        b"Sample resume text",
    )

    assert text == ""
    assert success is False
    assert error == "Unsupported file format."